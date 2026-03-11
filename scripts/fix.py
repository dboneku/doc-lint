#!/usr/bin/env python3
"""
doc-lint fixer: Apply auto-fixable formatting corrections to a .docx file.
Usage: python fix.py --file path/to/file.docx [--config .doc-lint.json] [--overwrite]
"""

import sys
import re
import json
import copy
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

MONOSPACE   = {'Courier New', 'Consolas', 'Monaco', 'Courier', 'Lucida Console'}
NON_DECIMAL = {'lowerRoman', 'upperRoman', 'lowerLetter', 'upperLetter'}
DEFAULT_CONFIG = {
    "rules": {
        "style-misuse":            {"enabled": True},
        "font-normalization":      {"enabled": True, "target-font": "Calibri"},
        "font-size-normalization": {"enabled": True, "sizes": {"h1": 20, "h2": 16, "h3": 14, "h4": 12, "body": 11}},
        "list-normalization":      {"enabled": True},
        "heading-level-skip":      {"enabled": True},
        "single-item-list":        {"enabled": True},
        "mixed-fonts":             {"enabled": True},
        "multiline-heading":       {"enabled": True},
        "numbered-heading-continuity": {"enabled": True},
    }
}

# Template for --init-config
_CONFIG_TEMPLATE = {
    "rules": {
        "consecutive-headings":        {"enabled": True,  "severity": "error"},
        "empty-section":               {"enabled": True,  "severity": "error"},
        "style-misuse":                {"enabled": True,  "severity": "warning"},
        "font-normalization":          {"enabled": True,  "severity": "warning", "target-font": "Calibri"},
        "font-size-normalization":     {"enabled": True,  "severity": "warning",
                                        "sizes": {"h1": 20, "h2": 16, "h3": 14, "h4": 12, "body": 11}},
        "list-normalization":          {"enabled": True,  "severity": "warning"},
        "heading-level-skip":          {"enabled": True,  "severity": "warning"},
        "single-item-list":            {"enabled": True,  "severity": "info"},
        "orphaned-bold":               {"enabled": True,  "severity": "info"},
        "mixed-fonts":                 {"enabled": True,  "severity": "info"},
        "multiline-heading":           {"enabled": True,  "severity": "info"},
        "numbered-heading-continuity": {"enabled": True,  "severity": "warning"},
        "template-compliance":         {"enabled": True,  "severity": "warning"},
        "naming-convention":           {"enabled": True,  "severity": "warning"},
        "style-policy":                {"enabled": True,  "severity": "warning"},
    }
}


def load_config(path):
    """Load config, merging user overrides on top of defaults.

    A rule value can be:
      - A dict:  {"enabled": true, "severity": "warning", ...extra...}
      - A string shorthand: "off" | "error" | "warning" | "info"
        "off" disables the rule; anything else sets enabled=True + severity.
    """
    cfg = json.loads(json.dumps(DEFAULT_CONFIG))
    if path and Path(path).exists():
        with open(path) as f:
            user = json.load(f)
        for rule, settings in user.get("rules", {}).items():
            if isinstance(settings, str):
                # shorthand: "off", "error", "warning", "info"
                if settings == "off":
                    cfg["rules"].setdefault(rule, {})["enabled"] = False
                else:
                    cfg["rules"].setdefault(rule, {}).update(
                        {"enabled": True, "severity": settings}
                    )
            elif isinstance(settings, dict):
                cfg["rules"].setdefault(rule, {}).update(settings)
    return cfg


def rule_enabled(cfg, rule):
    return cfg["rules"].get(rule, {}).get("enabled", True)


def get_para_size(para):
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    try:
        sz = para.style.font.size
        if sz:
            return sz.pt
    except Exception:
        pass
    return None


def get_numpr(pPr):
    if pPr is None:
        return None, None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None, None
    ilvl_el  = numPr.find(qn('w:ilvl'))
    numid_el = numPr.find(qn('w:numId'))
    numid = numid_el.get(qn('w:val')) if numid_el is not None else None
    return numid, int(ilvl_el.get(qn('w:val'))) if ilvl_el is not None else 0


def fix_style_misuse(doc, cfg, applied, changes):
    """W003 — Reclassify heading-styled body paragraphs as Normal."""
    if not rule_enabled(cfg, 'style-misuse'):
        return
    count = 0
    for para in doc.paragraphs:
        style = para.style.name
        size  = get_para_size(para)
        thresholds = {'Heading 1': 12, 'Heading 2': 10, 'Heading 3': 9}
        for h, thresh in thresholds.items():
            if h in style and size and size <= thresh:
                changes.append(('W003', para.text[:80], f'[style: {style} → Normal]'))
                para.style = doc.styles['Normal']
                count += 1
                break
    if count:
        applied.append(f"W003: Reclassified {count} misused heading paragraph(s) as Normal")


def fix_font_normalization(doc, cfg, applied, changes):
    """W004 — Normalize body text font family."""
    if not rule_enabled(cfg, 'font-normalization'):
        return
    target = cfg["rules"]["font-normalization"].get("target-font", "Calibri")
    count = 0
    for para in doc.paragraphs:
        if para.style.name not in ('Normal', 'Normal (Web)', 'Default Paragraph Style'):
            continue
        for run in para.runs:
            if run.font.name and run.font.name not in MONOSPACE and run.font.name != target:
                changes.append(('W004', para.text[:80],
                                 f'[font: "{run.font.name}" → "{target}"]'))
                run.font.name = target
                count += 1
    if count:
        applied.append(f"W004: Normalized {count} run(s) to {target}")


def fix_font_size(doc, cfg, applied, changes):
    """W005 — Normalize font sizes to standard scale."""
    if not rule_enabled(cfg, 'font-size-normalization'):
        return
    sizes = cfg["rules"]["font-size-normalization"].get("sizes", {})
    size_map = {
        'Heading 1': sizes.get('h1', 16),
        'Heading 2': sizes.get('h2', 14),
        'Heading 3': sizes.get('h3', 12),
        'Heading 4': sizes.get('h4', 12),
        'Normal':    sizes.get('body', 11),
        'Normal (Web)': sizes.get('body', 11),
    }
    count = 0
    for para in doc.paragraphs:
        target = next((v for k, v in size_map.items() if k in para.style.name), None)
        if target:
            for run in para.runs:
                if run.font.size and run.font.size.pt != target:
                    changes.append(('W005', para.text[:80],
                                    f'[size: {run.font.size.pt}pt → {target}pt]'))
                    run.font.size = Pt(target)
                    count += 1
    if count:
        applied.append(f"W005: Normalized {count} run(s) to standard font sizes")


def fix_list_normalization(doc, cfg, applied, changes):
    """W006 — Convert Roman/alphabetic ordered lists to decimal."""
    if not rule_enabled(cfg, 'list-normalization'):
        return
    count = 0
    try:
        numbering_el = doc.part.numbering_part._element
        for lvl in numbering_el.iter(qn('w:lvl')):
            el = lvl.find(qn('w:numFmt'))
            if el is not None and el.get(qn('w:val')) in NON_DECIMAL:
                old_fmt = el.get(qn('w:val'))
                changes.append(('W006', f'[list format: {old_fmt} → decimal]', ''))
                el.set(qn('w:val'), 'decimal')
                count += 1
    except Exception:
        pass
    if count:
        applied.append(f"W006: Normalized {count} list level(s) to decimal numbering")


def fix_heading_level_skip(doc, cfg, applied, changes):
    """W007 — Demote heading level skips to maintain sequential levels."""
    if not rule_enabled(cfg, 'heading-level-skip'):
        return
    current = 0
    count = 0
    for para in doc.paragraphs:
        style = para.style.name
        level = None
        for i in range(1, 7):
            if f'Heading {i}' in style:
                level = i
                break
        if level is None:
            current = 0
            continue
        if current > 0 and level > current + 1:
            new_level = current + 1
            try:
                changes.append(('W007', para.text[:80],
                                 f'[heading: H{level} → H{new_level}]'))
                para.style = doc.styles[f'Heading {new_level}']
                count += 1
                level = new_level
            except Exception:
                pass
        current = level
    if count:
        applied.append(f"W007: Fixed {count} heading level skip(s)")


def fix_single_item_lists(doc, cfg, applied, changes):
    """I008 — Convert single-item lists to plain paragraphs."""
    if not rule_enabled(cfg, 'single-item-list'):
        return
    groups = {}
    paras  = list(doc.paragraphs)
    for para in paras:
        pPr = para._element.find(qn('w:pPr'))
        numid, _ = get_numpr(pPr)
        if numid and numid != '0':
            groups.setdefault(numid, []).append(para)

    count = 0
    for numid, group in groups.items():
        if len(group) == 1:
            para = group[0]
            pPr  = para._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    changes.append(('I008', para.text[:80], '[list → paragraph]'))
                    pPr.remove(numPr)
            para.style = doc.styles['Normal']
            count += 1
    if count:
        applied.append(f"I008: Converted {count} single-item list(s) to paragraph(s)")


def fix_multiline_headings(doc, cfg, applied, changes):
    """I011 — Split multiline heading paragraphs at <w:br> soft breaks."""
    if not rule_enabled(cfg, 'multiline-heading'):
        return
    body  = doc.element.body
    count = 0
    for para_el in list(body):
        tag = para_el.tag.split('}')[-1] if '}' in para_el.tag else para_el.tag
        if tag != 'p':
            continue
        # Find runs with <w:br>
        split_run = None
        split_br  = None
        for r in para_el.findall(qn('w:r')):
            for br in r.findall(qn('w:br')):
                split_run = r
                split_br  = br
                break
            if split_run is not None:
                break
        if split_run is None:
            continue

        # Record text before the split for the diff
        full_text = ''.join(
            r.text for r in para_el.findall(qn('w:r'))
            if hasattr(r, 'text') and r.text
        )
        changes.append(('I011', full_text[:80], '[multiline heading split]'))

        # Remove the <w:br>
        split_run.remove(split_br)

        # Build new paragraph for content after the break
        new_para = OxmlElement('w:p')
        pPr = para_el.find(qn('w:pPr'))
        if pPr is not None:
            new_pPr = copy.deepcopy(pPr)
            pStyle  = new_pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                pStyle.set(qn('w:val'), 'Normal')
            new_para.insert(0, new_pPr)

        # Move runs after the split run to new paragraph
        after = False
        for child in list(para_el):
            child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if child == split_run:
                after = True
                continue
            if after and child_tag == 'r':
                para_el.remove(child)
                new_para.append(child)

        para_el.addnext(new_para)
        count += 1

    if count:
        applied.append(f"I011: Split {count} multiline heading paragraph(s) at line breaks")


def fix_numbered_headings(doc, cfg, applied, changes):
    """W012 — Strip numeric number prefixes from headings (e.g. '1. Purpose' → 'Purpose')."""
    if not rule_enabled(cfg, 'numbered-heading-continuity'):
        return
    numbered_pat = re.compile(r'^\ *\d+\.\s+')
    count = 0
    for para in doc.paragraphs:
        style = para.style.name
        hlevel = None
        for i in range(1, 7):
            if f'Heading {i}' in style:
                hlevel = i
                break
        if hlevel is None:
            continue
        m = numbered_pat.match(para.text.strip())
        if not m:
            continue
        prefix = m.group(0).lstrip()
        if para.runs and para.runs[0].text.lstrip().startswith(prefix.lstrip()):
            before = para.runs[0].text[:80]
            actual_prefix = para.runs[0].text[:len(para.runs[0].text) - len(para.runs[0].text.lstrip()) + len(prefix)]
            para.runs[0].text = para.runs[0].text[len(actual_prefix):]
            changes.append(('W012', before, para.runs[0].text[:80]))
            count += 1
    if count:
        applied.append(f"W012: Stripped number prefix from {count} heading(s)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Auto-fix formatting issues in a .docx file")
    parser.add_argument("--file",        required=False, help="Path to .docx file")
    parser.add_argument("--config",      default=".doc-lint.json", help="Config file path")
    parser.add_argument("--overwrite",   action="store_true", help="Overwrite original file")
    parser.add_argument("--init-config", action="store_true",
                        help="Write a template .doc-lint.json to the current directory and exit")
    args = parser.parse_args()

    if args.init_config:
        dest = Path(".doc-lint.json")
        if dest.exists():
            print(f"ERROR: {dest} already exists. Remove it first if you want to regenerate.")
            sys.exit(1)
        with open(dest, 'w') as f:
            json.dump(_CONFIG_TEMPLATE, f, indent=2)
            f.write('\n')
        print(f"Created {dest} — edit rules as needed.")
        print("  Set a rule to \"off\" to disable it, or \"error\"/\"warning\"/\"info\" to change severity.")
        return

    if not args.file:
        parser.error("--file is required unless --init-config is used")

    path = Path(args.file)
    if not path.exists():
        print(f"ERROR: File not found: {path}")
        sys.exit(1)

    cfg     = load_config(args.config)
    doc     = Document(str(path))
    applied = []
    changes = []  # list of (code, before_text, after_text)

    fix_style_misuse(doc, cfg, applied, changes)
    fix_font_normalization(doc, cfg, applied, changes)
    fix_font_size(doc, cfg, applied, changes)
    fix_list_normalization(doc, cfg, applied, changes)
    fix_heading_level_skip(doc, cfg, applied, changes)
    fix_single_item_lists(doc, cfg, applied, changes)
    fix_multiline_headings(doc, cfg, applied, changes)
    fix_numbered_headings(doc, cfg, applied, changes)

    out = path if args.overwrite else path.with_suffix('').parent / (path.stem + '.fixed.docx')
    doc.save(str(out))

    print(f"\nFixed: {out}")
    if applied:
        for a in applied:
            print(f"  ✓  {a}")
        if changes:
            # Before/after diff table
            col_w = 42
            print(f"\n  {'CODE':<6}  {'BEFORE':<{col_w}}  AFTER")
            print(f"  {'\u2500'*6}  {'\u2500'*col_w}  {'\u2500'*col_w}")
            for code, before, after in changes:
                b = (before[:col_w - 1] + '\u2026') if len(before) > col_w else before
                a = (after[:col_w - 1]  + '\u2026') if len(after)  > col_w else after
                print(f"  {code:<6}  {b:<{col_w}}  {a}")
    else:
        print("  No auto-fixable issues found.")
    print(f"\nRun /doc-lint:check {out} to verify.")


if __name__ == "__main__":
    main()
